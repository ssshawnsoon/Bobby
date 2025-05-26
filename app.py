import streamlit as st
from jamaibase import JamAI, protocol as p
from docx import Document
from io import BytesIO
import random
import string
from PyPDF2 import PdfReader
from tempfile import NamedTemporaryFile
from PIL import Image
import base64
import streamlit.components.v1 as components
import os

# Initialize JamAI
jamai = JamAI(api_key="jamai_pat_c4e44c670523685e19b2dee52f30119f795e26b6e0c2e65b", project_id="proj_1de330b27bce3d618b21379f")

def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"

st.set_page_config(page_title="Bobby.ai", page_icon="📝")
st.title("🌟 Bobby the Great - Your AI Assistant for Job Matching")

# Custom CSS
st.markdown(
    """
    <style>
    body {
        background-color: #1e1e1e;
        color: #f0f0f0;
    }
    .generated-output {
        background-color: #444;
        padding: 15px;
        border-radius: 10px;
        margin-top: 20px;
        box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.5);
        color: #f0f0f0;
    }
    .generated-output h4 {
        color: #FFA500;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Upload multiple CVs ---
with st.container():
    st.header("📄 Upload CVs for Comparison")

    uploaded_cvs = st.file_uploader("Upload CVs", type="pdf", accept_multiple_files=True)

    if st.button("🔍 Analyze & Choose Best Candidate"):
        if uploaded_cvs and 2 <= len(uploaded_cvs) <= 5:
            try:
                cv_texts = [extract_text_from_pdf(pdf) for pdf in uploaded_cvs]
                data_dict = {f"cv{i + 1}": text for i, text in enumerate(cv_texts)}

                result = jamai.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="TalentRank",
                        data=[data_dict],
                        stream=False
                    )
                )

                if result.rows:
                    best_candidate = result.rows[0].columns.get("Best_Candidate").text
                    st.success("✅ Best Candidate Identified")
                    st.markdown(f"**🏆 Best Candidate Summary:**\n\n{best_candidate}")

                    # Downloadable Report
                    doc = Document()
                    doc.add_heading("TalentRank - Best Candidate Report", level=1)
                    doc.add_paragraph(best_candidate)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📥 Download Report",
                        data=buffer,
                        file_name=generate_random_filename(),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("❌ No response from TalentRank.")
            except Exception as e:
                st.error(f"⚠️ Error during analysis: {e}")
        else:
            st.warning("⚠️ Please upload between 2 to 5 CVs.")

# CV Upload and Job Description Section
with st.container():
    st.header("📄 Interview Expert")
    st.write("Upload CV and Provide Job Description")
    cv_pdf = st.file_uploader("Upload CV (PDF format)", type="pdf", key="cv_upload_interview")
    job_description = st.text_area("✍️ Enter Job Description", key="job_desc_interview")

    # Button right here, right after inputs
    if st.button("🚀 Generate Interview Questions", use_container_width=True):
        if cv_pdf and job_description:
            try:
                cv_text = extract_text_from_pdf(cv_pdf)
                resume_completion = jamai.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="Interview_Question",
                        data=[{"cv": cv_text, "job_description": job_description}],
                        stream=False
                    )
                )

                if resume_completion.rows:
                    row = resume_completion.rows[0].columns
                    summary = row.get("summary")
                    generate_questions = row.get("generate_questions")

                    st.subheader("✨ Generated Output")
                    st.markdown(
                        f"""
                        <div class="generated-output">
                            <h4>📝 Summary:</h4> <p>{summary.text if summary else 'N/A'}</p>
                            <h4>💼 Interview Questions:</h4> <p>{generate_questions.text if generate_questions else 'N/A'}</p>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                    # Download report
                    with st.container():
                        st.subheader("📥 Download Final Report")
                        doc = Document()
                        doc.add_heading("Executive Report", level=1)
                        doc.add_heading("Interview Questions", level=2)
                        doc.add_paragraph(generate_questions.text if generate_questions else 'N/A')
                        doc.add_heading("Summary", level=2)
                        doc.add_paragraph(summary.text if summary else 'N/A')

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="📄 Download Final Report as .docx",
                            data=buffer,
                            file_name=generate_random_filename(),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("⚠️ Failed to process CV and job description.")
            except Exception as e:
                st.error(f"❌ Error during resume analysis: {e}")
        else:
            st.warning("⚠️ Please upload a CV and enter a job description.")

# MBTI Personality Analyzer Section
with st.container():
    st.header("🧠 MBTI Personality Analyzer")
    st.write("Answer the essay question below to generate your MBTI personality insights.")
    essay_input = st.text_area(
        "📝 Describe a time when you had to choose between following a logical solution and honoring your personal or team values. What did you choose and why? What was the outcome?"
    )

    # MBTI button right here, right after essay input
    if st.button("🧠 Analyze MBTI Essay", use_container_width=True):
        if essay_input.strip():
            try:
                mbti_completion = jamai.add_table_rows(
                    "action",
                    p.RowAddRequest(
                        table_id="mbti",
                        data=[{"essay": essay_input}],
                        stream=False
                    )
                )

                if mbti_completion.rows:
                    row = mbti_completion.rows[0].columns
                    MBTI_Report = row.get("MBTI_Report")

                    if MBTI_Report and MBTI_Report.text.strip():
                        formatted_text = MBTI_Report.text.replace('\n', '<br>')

                        st.subheader("✨ Generated Output")
                        st.markdown(
                            f"""
                            <div class="generated-output">
                                <h4>🧬 MBTI Personality Report:</h4>
                                <p>{formatted_text}</p>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                        # Download MBTI report
                        with st.container():
                            st.subheader("📥 Download MBTI Report")
                            doc = Document()
                            doc.add_heading("MBTI Personality Report", level=1)
                            doc.add_paragraph(MBTI_Report.text)

                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="📄 Download MBTI Report as .docx",
                                data=buffer,
                                file_name=generate_random_filename(),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    else:
                        st.warning("⚠️ MBTI report is missing or empty.")
                else:
                    st.warning("⚠️ MBTI analysis returned no rows.")
            except Exception as e:
                st.error(f"❌ Error during MBTI analysis: {e}")
        else:
            st.warning("⚠️ Please enter an essay to analyze.")

    # Fortune Face Analyzer
    st.header("🧞‍♀️ Fortune Face Reader ")
    st.markdown('<div class="title">The mystery of a face is its hidden tales ✦˖ ࣪𖤐₊˚ೄ</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader("Upload your beautiful face ✧˖° for fortune analysis", type=["jpg", "jpeg", "png"])
    analyze_button = st.button("🔍 Analyze Beautiful Face", use_container_width=True)

    if uploaded_file and analyze_button:
        st.write("Filename:", uploaded_file.name)

        # Convert uploaded file to JPEG
        with NamedTemporaryFile(delete=False, suffix=".jpeg") as temp_file:
            try:
                image = Image.open(uploaded_file).convert("RGB")
                image.save(temp_file, format="JPEG")
                temp_file_path = temp_file.name
            except Exception as e:
                st.error(f"❌ Image processing error: {e}")
                st.stop()

        # Upload image to JamAI
        try:
            upload_response = jamai.file.upload_file(temp_file_path)
        except Exception as e:
            st.error(f"❌ Failed to upload image: {e}")
            st.stop()

        # Analyze the face
        try:
            completion = jamai.table.add_table_rows(
                "action",
                p.RowAddRequest(
                    table_id="Fortune_Face",
                    data=[dict(Beautiful_Face=upload_response.uri)],
                    stream=False,
                ),
            )
        except Exception as e:
            st.error(f"❌ Analysis request failed: {e}")
            st.stop()

        # Display image preview
        image_bytes = uploaded_file.getvalue()
        encoded_image = base64.b64encode(image_bytes).decode()
        st.image(f"data:image/jpeg;base64,{encoded_image}", width=300, caption="Uploaded Image")

        if completion.rows:
            st.success("Image successfully analyzed!")
        else:
            st.error("Analysis failed. Please try again.")

        # Fetch result from the table
        rows = jamai.table.list_table_rows("action", "Fortune_Face")
        if rows.items:
            row = rows.items[0]

            face_attributes = {
                "Face Shape": row.get("Faceshape", {}).get("value", "N/A"),
                "Forehead": row.get("Forehead", {}).get("value", "N/A"),
                "Eyebrows": row.get("Eyebrows", {}).get("value", "N/A"),
                "Eyes": row.get("Eyes", {}).get("value", "N/A"),
                "Nose": row.get("Nose", {}).get("value", "N/A"),
                "Mouth & Lips": row.get("Mouth_and_Lips", {}).get("value", "N/A"),
                "Chin & Jawline": row.get("Chin_and_Jawline", {}).get("value", "N/A"),
                "Face Fortune Report": row.get("Face_Fortune_Report", {}).get("value", "N/A")
            }

            st.subheader("🔮 Fortune Face Analysis")
            for key, value in face_attributes.items():
                st.markdown(f"**{key}**: {value}")

            # Downloadable .docx report
            with st.container():
                st.subheader("📥 Download Fortune Face Report")
                doc = Document()
                doc.add_heading("Fortune Face Report", level=1)

                for key, value in face_attributes.items():
                    doc.add_heading(key, level=2)
                    doc.add_paragraph(value)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📄 Download Face Report as .docx",
                    data=buffer,
                    file_name=generate_random_filename(),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("⚠️ No results yet. Please try again after a moment.")

